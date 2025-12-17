from io import BytesIO
from docx import Document

def generate_word(company, heading, date, insights):
    doc = Document()

    doc.add_heading(company, 0)
    doc.add_heading(heading, level=1)
    doc.add_paragraph(f"Generated on: {date}")
    doc.add_page_break()

    for i in range(0, len(insights), 2):
        for j in range(2):
            if i + j < len(insights):
                doc.add_heading(f"Graph {i+j+1}", level=2)
                doc.add_paragraph("Graph Placeholder")
                doc.add_paragraph(insights[i+j])

        doc.add_page_break()

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
